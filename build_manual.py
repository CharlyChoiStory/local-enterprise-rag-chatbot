"""
MANUAL.docx 생성 스크립트
python-docx를 사용해 교육용 매뉴얼을 Word 파일로 만든다.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── 색상 팔레트 ───────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1E, 0x3A, 0x5F)   # 네이비 (제목)
BLUE_MID    = RGBColor(0x11, 0x64, 0xA3)   # 파랑 (소제목)
BLUE_LIGHT  = RGBColor(0xDE, 0xEB, 0xF7)   # 연파랑 (표 헤더)
GREEN_DARK  = RGBColor(0x1F, 0x6E, 0x43)   # 초록 (완료/성공)
GREEN_LIGHT = RGBColor(0xE2, 0xEF, 0xDA)   # 연초록 (표 배경)
ORANGE      = RGBColor(0xC5, 0x5A, 0x11)   # 주황 (경고)
ORANGE_LIGHT= RGBColor(0xFC, 0xE4, 0xD6)   # 연주황 (경고 배경)
GRAY_DARK   = RGBColor(0x40, 0x40, 0x40)   # 어두운 회색 (본문)
GRAY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)   # 연회색 (코드 배경)
GRAY_BOX    = RGBColor(0xD9, 0xD9, 0xD9)   # 회색 (박스)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


# ══════════════════════════════════════════════════════════════
#  헬퍼 함수
# ══════════════════════════════════════════════════════════════

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_table_border(table, size=6, color="1164A3"):
    """표 전체에 테두리 적용."""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def heading1(text, emoji=""):
    full = f"{emoji}  {text}" if emoji else text
    p = doc.add_heading(full, level=1)
    run = p.runs[0]
    run.font.color.rgb = BLUE_DARK
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # 하단 선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:color"), "1E3A5F")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def heading2(text, emoji=""):
    full = f"{emoji}  {text}" if emoji else text
    p = doc.add_heading(full, level=2)
    run = p.runs[0]
    run.font.color.rgb = BLUE_MID
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def heading3(text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = GRAY_DARK
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p


def body(text, bold=False, color=None, size=10.5, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color or GRAY_DARK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def code_block(lines: list[str]):
    """회색 배경 코드 블록."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    # 배경
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        if i < len(lines) - 1:
            run.add_break()
    return p


def info_box(lines: list[str], bg=BLUE_LIGHT, border_color="1164A3"):
    """강조 박스."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_table_border(tbl, size=8, color=border_color)
    cell.width = Cm(16)
    for line in lines:
        p   = cell.add_paragraph(line)
        run = p.runs[0] if p.runs else p.add_run(line)
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
    # 첫 번째 빈 단락 제거
    if cell.paragraphs and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def make_table(headers, rows,
               header_bg=BLUE_DARK, header_color=WHITE,
               alt_bg=BLUE_LIGHT, border_color="1164A3"):
    col_n = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=col_n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(tbl, color=border_color)

    # 헤더
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = header_color

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        bg = alt_bg if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def divider():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), "BFBFBF")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def bullet(text, level=0, emoji="•"):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run(f"{emoji}  {text}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY_DARK
    return p


def step_box(number, title, content_lines):
    """번호가 있는 단계 박스."""
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(tbl, color="1164A3")
    # 번호 칸
    nc = tbl.rows[0].cells[0]
    set_cell_bg(nc, BLUE_DARK)
    nc.width = Cm(1.5)
    np_ = nc.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr  = np_.add_run(str(number))
    nr.font.size  = Pt(20)
    nr.font.bold  = True
    nr.font.color.rgb = WHITE
    # 내용 칸
    cc = tbl.rows[0].cells[1]
    set_cell_bg(cc, BLUE_LIGHT)
    cp = cc.paragraphs[0]
    cr = cp.add_run(title)
    cr.font.bold = True
    cr.font.size = Pt(11)
    cr.font.color.rgb = BLUE_DARK
    for line in content_lines:
        p2  = cc.add_paragraph(line)
        p2r = p2.runs[0] if p2.runs else p2.add_run(line)
        p2r.font.size = Pt(9.5)
        p2r.font.color.rgb = GRAY_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════
#  표지
# ══════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(60)
r = p_title.add_run("📘  로컬 RAG 기업 챗봇 POC")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = BLUE_DARK

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("완전 매뉴얼 — AI 활용 교육 참가자용")
r2.font.size  = Pt(16)
r2.font.color.rgb = BLUE_MID

doc.add_paragraph()

info_box([
    "🏢  대상 독자   :  AI 활용 교육 참가자 (중장년 비개발직군)",
    "📅  작성 기준   :  2026-05-11",
    "🔗  GitHub      :  https://github.com/CharlyChoiStory/local-enterprise-rag-chatbot",
    "🛡️  보안 원칙   :  외부 AI API 없음 · 완전 로컬(On-Premise) 운영",
], bg=BLUE_LIGHT)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  목차 (수동)
# ══════════════════════════════════════════════════════════════

heading1("목차", "📋")

toc_items = [
    ("1장", "이 프로젝트는 무엇인가?"),
    ("2장", "RAG란 무엇인가? — 쉬운 비유로 이해하기"),
    ("3장", "왜 기업 보안이 중요한가?"),
    ("4장", "보안 원칙: 완전 내부 설치형(On-Premise) 개발"),
    ("5장", "전체 시스템 구조 한눈에 보기"),
    ("6장", "개발에 사용된 AI 도구들"),
    ("7장", "단계별 개발 과정 (상세)"),
    ("8장", "실제 실행 화면 안내"),
    ("9장", "성능 테스트 결과"),
    ("10장", "다음 확장 방향 및 용어 정리"),
]
make_table(
    ["장", "내용"],
    toc_items,
    header_bg=BLUE_DARK,
    alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  1장: 프로젝트 소개
# ══════════════════════════════════════════════════════════════

heading1("이 프로젝트는 무엇인가?", "1️⃣")
body(
    "회사 내부 문서(PDF·Word·PPT·TXT)를 AI가 읽고, "
    "임직원이 자연어로 질문하면 문서 근거를 바탕으로 "
    "자동 답변해 주는 챗봇을 만들었습니다.",
    size=11,
)

heading2("해결하는 문제", "🎯")
make_table(
    ["기존 방식 (불편)", "이 챗봇 방식 (편리)"],
    [
        ("담당자에게 전화·메일 문의", "챗봇에 바로 질문"),
        ("여러 폴더를 직접 뒤져서 문서 찾기", "자연어 질문 → 자동 검색"),
        ("같은 질문 반복 응대 (담당자 피로)", "24시간 자동 응답"),
        ("퍼블릭 AI에 회사 문서 업로드 (보안 위험)", "회사 내부망에서만 동작 (안전)"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("대화 예시", "💬")
info_box([
    "👤  직원   :  \"연차는 며칠 전까지 신청해야 해?\"",
    "",
    "🤖  챗봇   :  최소 3영업일 전까지 HR포털에서 신청하셔야 합니다. [1]",
    "              단, 3일 이상 연속 사용 시에는 5영업일 전 신청이 필요합니다. [2]",
    "",
    "📎  출처   :  [1] 연차휴가규정.txt · 단락 3 · 유사도 0.73",
    "              [2] 01_연차_휴가_운영_가이드.pdf · page 1 · 유사도 0.70",
], bg=GREEN_LIGHT, border_color="1F6E43")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  2장: RAG란 무엇인가?
# ══════════════════════════════════════════════════════════════

heading1("RAG란 무엇인가? — 쉬운 비유로 이해하기", "2️⃣")

heading2("RAG 단어 뜻 풀이", "🔤")
make_table(
    ["알파벳", "영어", "한국어", "역할"],
    [
        ("R", "Retrieval", "검색", "관련 문서 조각을 DB에서 찾아오는 것"),
        ("A", "Augmented", "보강", "찾은 내용으로 AI 답변 재료를 보충하는 것"),
        ("G", "Generation", "생성", "AI가 최종 자연어 답변을 만들어 내는 것"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("도서관 사서 비유", "📚")
info_box([
    "━━━ 일반 AI 챗봇 (ChatGPT 같은 퍼블릭 AI) ━━━",
    "",
    "❓ 질문 : \"우리 회사 연차 규정이 어떻게 돼?\"",
    "🤖 답변 : \"일반적으로 1년에 15일 정도입니다...\" ← 인터넷 지식으로 추측",
    "          (우리 회사 규정을 전혀 모름!)",
], bg=ORANGE_LIGHT, border_color="C55A11")

info_box([
    "━━━ RAG 챗봇 (이 시스템) ━━━",
    "",
    "❓ 질문 : \"우리 회사 연차 규정이 어떻게 돼?\"",
    "📚 1단계 : 회사 문서 창고에서 관련 문서를 자동으로 찾는다",
    "📄 2단계 : 「연차휴가규정.txt」 3번 단락 발견!",
    "🤖 3단계 : 그 내용을 바탕으로 정확한 답변 생성",
    "✅ 답변 : \"최소 3영업일 전까지 HR포털에서 신청하셔야 합니다.\"",
    "📎 출처 : 연차휴가규정.txt · 단락 3",
], bg=GREEN_LIGHT, border_color="1F6E43")

heading2("RAG 작동 순서 (단계별)", "🔄")
steps_rag = [
    ("STEP 1", "질문 입력",
     "사용자가 자연어로 질문을 입력합니다.",
     "예) \"출장비 영수증은 어떻게 제출해?\""),
    ("STEP 2", "질문을 숫자(벡터)로 변환 — 임베딩",
     "AI가 질문의 의미를 숫자 목록으로 변환합니다.",
     "\"출장비 영수증\" → [0.23, 0.87, 0.41, ...]"),
    ("STEP 3", "유사한 문서 조각 Top-5 검색",
     "미리 저장된 문서 DB에서 가장 관련성 높은 내용을 찾습니다.",
     "Chroma Vector DB에서 코사인 유사도로 검색"),
    ("STEP 4", "검색된 내용 + 질문을 AI에 전달",
     "찾은 문서 조각 5개를 묶어 AI(LLM)에 함께 전달합니다.",
     "\"다음 문서를 참고해서 답변해줘: [문서 내용]...\""),
    ("STEP 5", "로컬 AI가 답변 생성",
     "llama3.1:8b 모델이 문서 근거를 바탕으로 답변을 만듭니다.",
     "출처 번호 [1][2]를 붙여 신뢰성을 표시합니다."),
]
for s in steps_rag:
    tbl = doc.add_table(rows=1, cols=2)
    set_table_border(tbl)
    nc = tbl.rows[0].cells[0]
    set_cell_bg(nc, BLUE_DARK)
    nc.width = Cm(2.0)
    np_ = nc.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr  = np_.add_run(s[0])
    nr.font.size = Pt(9)
    nr.font.bold = True
    nr.font.color.rgb = WHITE
    cc = tbl.rows[0].cells[1]
    set_cell_bg(cc, BLUE_LIGHT)
    cp = cc.paragraphs[0]
    cr = cp.add_run(s[1])
    cr.font.bold = True
    cr.font.size = Pt(10.5)
    cr.font.color.rgb = BLUE_DARK
    for line in s[2:]:
        p2  = cc.add_paragraph(line)
        r2  = p2.runs[0] if p2.runs else p2.add_run(line)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRAY_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

heading2("전통 키워드 검색 vs RAG 비교", "🆚")
make_table(
    ["항목", "전통 키워드 검색", "RAG 검색 (이 시스템)"],
    [
        ("검색 방식", "단어가 정확히 일치해야 함", "의미와 문맥이 비슷하면 검색"),
        ("질문 예시", "\"연차 신청\" 입력 필요", "\"휴가 얼마 전에 내야 해?\"도 검색됨"),
        ("결과 형태", "파일 목록 반환", "자연어 답변 반환"),
        ("출처 표시", "없음", "파일명 + 위치 + 유사도 표시"),
        ("정확도", "키워드가 없으면 실패", "의미가 비슷하면 성공"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  3장: 기업 보안 필요성
# ══════════════════════════════════════════════════════════════

heading1("왜 기업 보안이 중요한가?", "3️⃣")

heading2("퍼블릭 AI 사용 시 발생하는 보안 위험", "🚨")
info_box([
    "직원이 ChatGPT에 회사 문서 내용을 붙여넣을 때 일어나는 일",
    "",
    "직원 PC  →  인터넷 경유  →  AI 회사 서버 (해외)",
    "",
    "⚠️  질문 내용이 외부 서버에 저장됨",
    "⚠️  AI 학습 데이터로 활용될 수 있음",
    "⚠️  경쟁사에 정보가 간접 유출될 수 있음",
    "⚠️  개인정보보호법 위반 가능성 (GDPR 등)",
], bg=ORANGE_LIGHT, border_color="C55A11")

heading2("기업에서 실제 발생한 사례", "📰")
make_table(
    ["기업/기관", "사건 내용", "결과"],
    [
        ("삼성전자 (2023)", "반도체 소스코드를 ChatGPT에 입력", "기밀 유출 우려 → 사내 ChatGPT 사용 전면 금지"),
        ("미국 대형 로펌", "법률 문서를 AI 서비스에 업로드", "의뢰인 개인정보 유출 문제 → 소송 위기"),
        ("국내 금융기관 다수", "고객 데이터를 AI에 입력", "내부 규정 신설: 고객정보 AI 입력 금지"),
    ],
    header_bg=RGBColor(0xC5, 0x5A, 0x11), header_color=WHITE, alt_bg=ORANGE_LIGHT,
)

heading2("기업이 지켜야 할 보안 원칙", "🔐")
make_table(
    ["보안 원칙", "설명"],
    [
        ("데이터 주권", "회사 데이터는 회사가 통제할 수 있는 곳에만 저장"),
        ("전송 최소화", "불필요하게 외부로 데이터를 보내지 않음"),
        ("접근 제어", "필요한 사람만 필요한 데이터에 접근"),
        ("감사 가능성", "누가 언제 무엇을 조회했는지 기록"),
        ("컴플라이언스", "개인정보보호법, GDPR 등 법적 요구사항 준수"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  4장: On-Premise 보안 아키텍처
# ══════════════════════════════════════════════════════════════

heading1("보안 원칙: 완전 내부 설치형(On-Premise) 개발", "4️⃣")

heading2("이 시스템의 보안 구조", "🔒")
info_box([
    "🏢  회사 내부 네트워크 경계",
    "",
    "  직원 PC(브라우저)  →  Streamlit 웹 UI  →  Python 앱",
    "                                               ↓",
    "                              Chroma Vector DB (로컬 저장)",
    "                                               ↓",
    "                              Ollama 로컬 LLM (내 PC에서 실행)",
    "                                               ↓",
    "                              회사 문서 폴더 (data/source_docs/)",
    "",
    "═══════════════════════════════════════════════",
    "🚫  인터넷 구간으로 나가는 데이터: 완전 없음",
    "🚫  외부 AI API 호출: 없음 (OpenAI·Anthropic·Google 금지)",
    "✅  모든 AI 처리: 회사 내부에서만 실행",
], bg=GREEN_LIGHT, border_color="1F6E43")

heading2("적용된 보안 규칙 체크리스트", "✅")
make_table(
    ["보안 항목", "적용", "구현 방법"],
    [
        ("퍼블릭 AI API 호출 금지", "✅ 적용", "Ollama 로컬 모델 사용"),
        ("외부 망으로 문서 전송 금지", "✅ 적용", "모든 처리 로컬 실행"),
        ("외부 AI API Key 없음", "✅ 적용", ".env에 외부 AI key 항목 없음"),
        ("회사 문서 로컬 보관", "✅ 적용", "data/source_docs/ 에만 저장"),
        ("Vector DB 로컬 보관", "✅ 적용", "data/chroma/ 에만 저장"),
        ("로그에 원문 전체 저장 금지", "✅ 적용", "파일명·chunk ID·점수만 기록"),
        (".env 파일 Git 업로드 금지", "✅ 적용", ".gitignore에 .env 포함"),
    ],
    header_bg=BLUE_DARK, alt_bg=GREEN_LIGHT,
)

heading2("퍼블릭 AI vs 이 시스템 비교", "🆚")
make_table(
    ["항목", "ChatGPT 등 퍼블릭 AI", "이 시스템 (로컬 On-Premise)"],
    [
        ("데이터 이동", "인터넷 경유 (위험)", "내부망만 (안전)"),
        ("문서 기밀성", "보장 안 됨", "완전 보장"),
        ("사용 비용", "월정액·종량제 유료", "무료 (오픈소스)"),
        ("인터넷 필요 여부", "필수", "불필요"),
        ("커스터마이징", "제한적", "완전 자유"),
        ("법적 컴플라이언스", "어려움", "쉬움 (내부 통제)"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  5장: 시스템 구조
# ══════════════════════════════════════════════════════════════

heading1("전체 시스템 구조 한눈에 보기", "5️⃣")

heading2("데이터 흐름 — 문서 준비 단계", "📥")
make_table(
    ["단계", "처리 내용", "사용 도구", "결과물"],
    [
        ("1. 문서 수집", "폴더 스캔 및 파일 목록 생성", "Python pathlib", "파일 경로 목록"),
        ("2. 텍스트 추출", "파일 형식별 텍스트 파싱", "PyMuPDF / python-docx\n/ python-pptx / charset-normalizer", "섹션별 텍스트"),
        ("3. 청킹", "1600자 단위 조각 분리\n(250자 겹침)", "src/chunker.py", "67개 chunk"),
        ("4. 임베딩", "텍스트를 숫자 벡터로 변환", "Ollama bge-m3 모델", "숫자 벡터 배열"),
        ("5. DB 저장", "벡터 + 메타데이터 저장", "Chroma Vector DB", "검색 가능한 DB"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("데이터 흐름 — 질문 응답 단계", "💬")
make_table(
    ["단계", "처리 내용", "사용 도구", "결과물"],
    [
        ("1. 질문 입력", "사용자가 자연어로 질문", "Streamlit UI", "질문 텍스트"),
        ("2. 질문 임베딩", "질문을 숫자 벡터로 변환", "Ollama bge-m3 모델", "쿼리 벡터"),
        ("3. 유사 문서 검색", "코사인 유사도로 Top-5 검색", "Chroma Vector DB", "관련 chunk 5개"),
        ("4. 컨텍스트 조립", "문서 조각 + 질문 묶기", "src/rag.py", "프롬프트 텍스트"),
        ("5. 답변 생성", "LLM이 문서 기반 답변 생성", "Ollama llama3.1:8b", "자연어 답변 + 출처"),
    ],
    header_bg=BLUE_DARK, alt_bg=GREEN_LIGHT,
)

heading2("소프트웨어 기술 스택", "📦")
make_table(
    ["역할", "사용 도구", "특징"],
    [
        ("화면 (UI)", "Streamlit", "Python으로 웹앱 빠르게 제작"),
        ("PDF 읽기", "PyMuPDF", "고속 PDF 텍스트 추출"),
        ("Word 읽기", "python-docx", ".docx 파일 파싱"),
        ("PPT 읽기", "python-pptx", ".pptx 슬라이드별 텍스트"),
        ("TXT 읽기", "charset-normalizer", "한글 인코딩 자동 감지"),
        ("로컬 AI 실행", "Ollama", "LLM·임베딩 모델 로컬 실행 플랫폼"),
        ("채팅 LLM", "llama3.1:8b (4.9GB)", "Meta 오픈소스 대형 언어 모델"),
        ("임베딩 모델", "bge-m3 (1.2GB)", "다국어 문장 의미 변환 모델"),
        ("벡터 DB", "Chroma", "코사인 유사도 기반 로컬 Vector DB"),
        ("설정 관리", "pydantic-settings", "환경변수 검증 및 타입 안전성"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)
info_box(["✅  모두 무료 오픈소스   |   ✅  외부 AI API 없음   |   ✅  인터넷 없이 동작"],
         bg=GREEN_LIGHT, border_color="1F6E43")
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  6장: AI 도구
# ══════════════════════════════════════════════════════════════

heading1("개발에 사용된 AI 도구들", "6️⃣")

heading2("AI 협업 생태계 구성도", "🤖")
make_table(
    ["도구", "역할", "활용 방식"],
    [
        ("📱  텔레그램 (Telegram)", "외부 원격 지시 채널",
         "개발자가 어디서든 텔레그램으로\nOpenClaw에 개발 지시 전달"),
        ("🦅  OpenClaw 에이전트", "전체 진행 컨설팅 및 원격 지시 해석",
         "텔레그램 지시 수신 → 분석 → 단계별 실행 계획\n기획·아키텍처·테스트 전략 수립"),
        ("💻  VSCode + Claude Code", "로컬 PC 코드 자동 생성 및 실행",
         "자연어 지시 → 코드 자동 작성\n디버깅·리팩토링·테스트 지원"),
        ("🐙  GitHub", "소스코드 버전 관리 및 공유",
         "Git 이력 관리 + 오픈소스 공개\n팀 협업 및 코드 백업"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

# ── 텔레그램 → OpenClaw → 로컬 시스템 흐름 도식 ──────────────
heading2("외부 원격 지시 흐름 도식: Telegram → OpenClaw → 로컬 시스템", "📡")
body("보안 원칙상 외부에서 로컬 시스템에 직접 접근하지 않고, "
     "텔레그램 → OpenClaw 에이전트 → Claude Code(로컬) 경로로만 개발 지시가 전달되었습니다.",
     size=10.5)

# 도식 (표로 구현)
def _remote_flow_diagram():
    rows_data = [
        # (배경색, 아이콘, 제목, 설명)
        (RGBColor(0x1E, 0x3A, 0x5F), "📱", "외부 환경 (개발자·관리자)",
         "외부 어디서든 스마트폰·PC로 접속 가능"),
        None,  # 화살표 행
        (RGBColor(0x17, 0x52, 0x8A), "🦅", "텔레그램 → OpenClaw 에이전트",
         "\"2단계: 파서 구현해줘. PDF·TXT·DOCX·PPTX 지원\" 등 자연어 지시"),
        None,
        (RGBColor(0x1F, 0x6E, 0x43), "🤖", "Claude Code  (VSCode + 로컬 PC)",
         "자연어 지시를 코드로 변환 · 자동 실행 · 오류 수정"),
        None,
        (RGBColor(0x40, 0x40, 0x70), "🏠", "로컬 RAG 챗봇 시스템",
         "Streamlit UI + Ollama LLM + Chroma DB  ← 모두 로컬(인터넷 없음)"),
        None,
        (RGBColor(0x17, 0x52, 0x8A), "🦅", "OpenClaw 에이전트 → 텔레그램 보고",
         "\"파서 완료. 30/30 성공. 다음 단계 진행하시겠습니까?\""),
        None,
        (RGBColor(0x1E, 0x3A, 0x5F), "📱", "외부 환경 (개발자·관리자)",
         "결과 확인 후 다음 단계 지시"),
    ]

    for row in rows_data:
        if row is None:
            # 화살표 행
            p_arr = doc.add_paragraph()
            p_arr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_arr = p_arr.add_run("▼")
            r_arr.font.size = Pt(14)
            r_arr.font.color.rgb = BLUE_MID
            p_arr.paragraph_format.space_before = Pt(0)
            p_arr.paragraph_format.space_after  = Pt(0)
        else:
            bg_color, icon, title, desc = row
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            set_table_border(tbl, size=6, color="1164A3")

            # 아이콘 셀
            ic = tbl.rows[0].cells[0]
            set_cell_bg(ic, bg_color)
            ic.width = Cm(1.4)
            ip = ic.paragraphs[0]
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ir = ip.add_run(icon)
            ir.font.size = Pt(18)

            # 내용 셀
            cc = tbl.rows[0].cells[1]
            set_cell_bg(cc, RGBColor(
                min(bg_color[0] + 80, 255),
                min(bg_color[1] + 80, 255),
                min(bg_color[2] + 80, 255),
            ))
            tp = cc.paragraphs[0]
            tr = tp.add_run(title)
            tr.font.bold = True
            tr.font.size = Pt(10.5)
            tr.font.color.rgb = WHITE
            dp = cc.add_paragraph(desc)
            dr = dp.runs[0] if dp.runs else dp.add_run(desc)
            dr.font.size = Pt(9)
            dr.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

            doc.add_paragraph().paragraph_format.space_after = Pt(0)

_remote_flow_diagram()
doc.add_paragraph().paragraph_format.space_after = Pt(6)

info_box([
    "🔐  보안 포인트",
    "",
    "  • 텔레그램은 지시 전달 채널일 뿐 — 회사 문서나 RAG 처리 결과는 텔레그램을 통해",
    "    오가지 않습니다.",
    "  • 로컬 RAG 시스템은 항상 내부망에서만 동작하며,",
    "    외부에서 직접 접근하는 포트나 URL은 존재하지 않습니다.",
    "  • OpenClaw 에이전트가 전달하는 것은 '개발 지시(텍스트)'뿐입니다.",
], bg=GREEN_LIGHT, border_color="1F6E43")

heading2("OpenClaw 원격 지시 — 단계별 예시", "📋")
make_table(
    ["단계", "위치", "실제 내용"],
    [
        ("1. 지시 전달", "외부 → 텔레그램 → OpenClaw",
         "\"2단계: 파서 구현. PDF·TXT·DOCX·PPTX 지원, 카테고리 자동 감지\""),
        ("2. 계획 수립", "OpenClaw 에이전트",
         "요구사항 정리 → 파일 구조 설계 → Claude Code에 구현 지시"),
        ("3. 코드 생성", "Claude Code (로컬 PC)",
         "src/parsers.py 자동 생성 → 30개 파일 테스트 실행"),
        ("4. 결과 보고", "Claude Code → OpenClaw → 텔레그램",
         "\"파서 완료. 30/30 성공. 오류 0건. 3단계 진행 가능\""),
        ("5. 다음 지시", "외부 → 텔레그램 → OpenClaw",
         "\"3단계 진행. 청킹 1600자·겹침 250자 기준으로\""),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("바이브코딩(Vibe Coding)이란?", "🎵")
info_box([
    "기존 개발 방식:",
    "  개발자가 코드를 한 줄씩 직접 타이핑 (수백 시간 소요)",
    "",
    "바이브코딩 방식:",
    "  사람이 \"무엇을 만들지\" 말로 설명하면",
    "  AI(Claude Code)가 코드를 대신 작성",
    "",
    "  👨‍💼 사람 (감독자·기획자, 텔레그램으로 지시)",
    "      \"파서를 구현해줘. PDF·TXT·DOCX·PPTX 지원하고...\"",
    "         ↓  OpenClaw 에이전트가 해석·전달",
    "  🤖 Claude Code (로컬 PC에서 실행)",
    "      → src/parsers.py 자동 생성",
    "      → 30개 파일 테스트 코드 실행",
    "      → 버그 자동 수정 후 보고",
    "",
    "마치 건축가가 설계도를 그리면 시공팀이 건물을 짓는 것과 같습니다.",
], bg=BLUE_LIGHT)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  7장: 단계별 개발 과정
# ══════════════════════════════════════════════════════════════

heading1("단계별 개발 과정 (상세)", "7️⃣")

heading2("전체 개발 타임라인", "🗓️")
make_table(
    ["단계", "내용", "완료 기준"],
    [
        ("0단계: 뼈대 생성", "폴더 구조·설정 파일 생성", "streamlit run app.py 실행 시 화면 표시"),
        ("1단계: 샘플 문서", "PDF 20개 + TXT 10개 준비", "30개 파일이 source_docs에 존재"),
        ("2단계: 문서 파서", "PDF·TXT·DOCX·PPTX 파서 구현", "파싱 성공률 100%"),
        ("3단계: 청킹", "1600자/250자 겹침 구현", "67개 chunk 생성 확인"),
        ("4단계: 임베딩+DB", "bge-m3 임베딩·Chroma 저장", "67개 벡터 DB 저장 완료"),
        ("5단계: RAG 파이프라인", "LLM 연결·답변 생성 구현", "샘플 질문 답변 확인"),
        ("6단계: Slack UI", "Streamlit Slack 스타일 UI", "인덱싱·채팅·출처 표시 동작"),
        ("7단계: GitHub 배포", "소스코드 GitHub 공개 및 버전 관리", "GitHub 리포지토리 접근 가능"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("0단계: 프로젝트 뼈대 생성", "🏗️")
body("폴더 구조와 설정 파일을 생성하고, Streamlit 빈 화면이 뜨는 상태를 목표로 합니다.")
code_block([
    "local-rag-poc/",
    "├── app.py                 ← 메인 웹앱",
    "├── requirements.txt       ← 필요 라이브러리 목록",
    "├── .env.example           ← 환경설정 예시 (외부 API key 없음)",
    "├── README.md              ← 설치·실행 가이드",
    "├── data/",
    "│   ├── source_docs/       ← 원본 회사 문서",
    "│   │   ├── hr/            ← 인사·연차 문서",
    "│   │   ├── finance/       ← 출장비·경비 문서",
    "│   │   ├── welfare/       ← 복지·경조사 문서",
    "│   │   ├── it/            ← IT장비·보안 문서",
    "│   │   └── policy/        ← 사내규정·매뉴얼",
    "│   └── chroma/            ← AI 검색용 벡터 DB (자동 생성)",
    "└── src/",
    "    ├── config.py          ← 설정값 관리",
    "    ├── parsers.py         ← 문서 읽기",
    "    ├── chunker.py         ← 문서 조각내기",
    "    ├── embeddings.py      ← 숫자로 변환",
    "    ├── vector_store.py    ← DB 저장·검색",
    "    ├── llm.py             ← AI 답변 생성",
    "    ├── rag.py             ← 전체 파이프라인",
    "    └── sample_questions.py ← 샘플 질문 목록",
])

heading2("2단계: 문서 파서 구현", "🔍")
make_table(
    ["파일 형식", "사용 라이브러리", "추출 단위", "비고"],
    [
        (".pdf", "PyMuPDF (fitz)", "페이지 단위", "page 1, page 2..."),
        (".txt", "charset-normalizer", "단락 단위", "한글 인코딩 자동 감지"),
        (".docx", "python-docx", "단락 단위", "Word 스타일 유지"),
        (".pptx", "python-pptx", "슬라이드 단위", "slide 1, slide 2..."),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("3단계: 청킹(Chunking) 원리", "✂️")
info_box([
    "청킹이란? 긴 문서를 AI가 소화할 수 있는 크기로 조각내는 것",
    "",
    "설정값:",
    "  • chunk 크기   : 1,600자",
    "  • overlap(겹침): 250자  ← 문단 경계에서 내용이 잘리지 않도록",
    "",
    "결과: 30개 문서 → 67개 chunk",
    "",
    "겹침(overlap)이 필요한 이유:",
    "  문서를 정확히 1600자씩 자르면 중요한 내용이 잘릴 수 있음.",
    "  앞 chunk의 끝 250자를 다음 chunk에도 포함시켜 문맥을 이어줌.",
], bg=BLUE_LIGHT)

heading2("4단계: 임베딩 — 텍스트를 숫자로", "🔢")
info_box([
    "임베딩이란? 텍스트의 '의미'를 AI가 계산할 수 있는 숫자 목록으로 변환하는 것",
    "",
    "\"연차는 3영업일 전까지\"   →  [0.23, 0.87, 0.41, 0.92, ...]",
    "\"휴가 신청 마감일\"         →  [0.21, 0.85, 0.44, 0.90, ...]  ← 숫자가 비슷! (의미 유사)",
    "\"오늘 점심 뭐 먹지?\"       →  [0.91, 0.12, 0.73, 0.05, ...]  ← 숫자가 다름 (의미 다름)",
    "",
    "사용 모델: bge-m3 (Ollama)",
    "저장 위치: data/chroma/ (Chroma Vector DB)",
    "저장 완료: 67개 chunk → 23초 소요",
], bg=BLUE_LIGHT)

heading2("5단계: RAG 답변 생성 흐름", "💬")
make_table(
    ["처리 단계", "상세 내용"],
    [
        ("① 질문 임베딩", "\"연차 신청 기한\" → [0.24, 0.86, 0.42, ...]"),
        ("② 유사 chunk 검색", "Chroma DB에서 코사인 유사도 Top-5 반환"),
        ("③ 근거 부족 판단", "유사도 0.30 미만이면 \"확인할 수 없습니다\" 출력"),
        ("④ 컨텍스트 조립", "찾은 문서 조각 5개를 합쳐 최대 4,000자 컨텍스트 구성"),
        ("⑤ LLM 호출", "llama3.1:8b에 [시스템 프롬프트 + 컨텍스트 + 질문] 전달"),
        ("⑥ 답변 반환", "출처 번호 [1][2]가 포함된 한국어 답변 생성"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("시스템 프롬프트 (AI에게 주는 지시)", "📜")
code_block([
    "너는 회사 내부 문서 기반 업무지원 챗봇이다.",
    "아래 제공된 문서 근거만 사용해서 한국어로 답변하라.",
    "근거에 없는 내용은 추측하지 말고",
    "  '제공된 문서에서 확인할 수 없습니다'라고 답하라.",
    "답변에는 [1], [2] 같은 출처 번호를 붙여라.",
    "연차, 출장비, 복지포인트, 회사 내규 질문은",
    "  실무자가 바로 이해할 수 있게 간결하게 답하라.",
])

heading2("6단계: Slack 스타일 UI", "🖥️")
make_table(
    ["UI 구성 요소", "위치", "기능"],
    [
        ("앱 이름·설명", "좌측 사이드바 상단", "\"사내 문서 챗봇\" · \"외부 AI 없음\""),
        ("설정 패널", "좌측 사이드바", "문서 폴더·모델명·Top-K 슬라이더 조정"),
        ("인덱싱 버튼", "좌측 사이드바", "문서 파싱→청킹→임베딩→DB저장 실행"),
        ("인덱싱 상태", "좌측 사이드바", "완료/미완 상태 표시 + chunk 수 확인"),
        ("샘플 질문 버튼", "좌측 사이드바 하단", "클릭 시 자동 질문 전송"),
        ("채팅 타임라인", "중앙 메인 영역", "Slack 스타일 말풍선 (파랑·회색)"),
        ("출처 박스", "봇 답변 하단", "파일명·위치·유사도·snippet 표시"),
        ("입력창", "하단 고정", "질문 입력 + 전송 버튼"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("7단계: GitHub 배포", "🚀")
body(
    "보안 원칙상 Streamlit + Ollama 기반 로컬 시스템은 외부 서버리스 플랫폼(Vercel 등)에 "
    "배포할 수 없습니다. 소스코드만 GitHub에 공개하여 버전 관리 및 재현성을 확보했습니다.",
    size=10.5,
)
info_box([
    "⚠️  Vercel·클라우드 배포가 불가능한 이유",
    "",
    "  1. Streamlit 앱은 지속적인 WebSocket 서버가 필요  →  Vercel 서버리스 구조 불가",
    "  2. Ollama 바이너리(5GB)를 서버리스 함수(50MB 제한)에 올릴 수 없음",
    "  3. Chroma Vector DB는 영구 파일시스템 필요  →  Vercel /tmp(512MB, 휘발성) 불가",
    "  4. 기업 내부 문서를 외부 클라우드에 올리는 것 자체가 보안 위반",
    "",
    "✅  해결 방법: GitHub에 소스코드 공개 → 사내 서버/개발자 PC에서 직접 실행",
], bg=ORANGE_LIGHT, border_color="C55A11")
make_table(
    ["단계", "명령어 / 내용"],
    [
        ("① .gitignore 설정", "보안 파일(.env)·가상환경(.venv)·Vector DB(data/chroma/) 제외"),
        ("② git init", "로컬 저장소 초기화"),
        ("③ git add & commit", "전체 소스코드 스테이징 및 커밋"),
        ("④ GitHub 리포지토리 생성", "gh repo create local-enterprise-rag-chatbot --public"),
        ("⑤ git push", "GitHub에 소스코드 업로드"),
        ("⑥ 실행 방법 공유", "README.md의 설치·실행 가이드를 팀원에게 전달"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  8장: 실행 방법
# ══════════════════════════════════════════════════════════════

heading1("실제 실행 화면 안내", "8️⃣")

heading2("처음 시작하는 순서", "🚀")

steps_run = [
    (1, "Ollama 설치",
     ["macOS 터미널에서:  brew install ollama",
      "✅ 완료 메시지: \"ollama was successfully installed\""]),
    (2, "AI 모델 다운로드",
     ["$ ollama pull bge-m3        ← 임베딩 모델 (1.2GB, 약 2-3분)",
      "$ ollama pull llama3.1:8b   ← 채팅 모델  (4.9GB, 약 5-10분)",
      "✅ 완료 메시지: \"success\""]),
    (3, "Python 환경 설정 및 앱 실행",
     ["$ cd local-rag-poc",
      "$ python3 -m venv .venv",
      "$ source .venv/bin/activate",
      "$ pip install -r requirements.txt",
      "$ streamlit run app.py",
      "✅ 브라우저 자동 열림: http://localhost:8501"]),
    (4, "문서 인덱싱",
     ["브라우저에서 왼쪽 사이드바 확인",
      "\"인덱싱 실행\" 버튼 클릭",
      "✅ \"완료! 67개 chunk 저장됨\" 메시지 확인"]),
    (5, "질문하기",
     ["하단 입력창에 질문 입력",
      "예) \"연차는 며칠 전까지 신청해야 해?\"",
      "→ 전송 버튼 클릭 → 답변 + 출처 확인"]),
]
for num, title, lines in steps_run:
    step_box(num, title, lines)

heading2("테스트 질문 모음", "🧪")
make_table(
    ["카테고리", "테스트 질문", "예상 출처 파일"],
    [
        ("HR (인사)", "연차는 며칠 전까지 신청해야 해?", "연차휴가규정.txt"),
        ("HR (인사)", "재택근무 신청은 어떻게 해?", "재택근무운영지침.txt"),
        ("재무", "출장비 영수증은 어떻게 제출해?", "02_출장비_정산_규정.pdf"),
        ("재무", "법인카드 사용 원칙이 뭐야?", "08_법인카드_사용_지침.pdf"),
        ("복지", "복지포인트 잔액은 어디서 확인해?", "복지포인트안내.txt"),
        ("복지", "경조사 지원금 신청 방법 알려줘", "경조사지원규정.txt"),
        ("IT", "신규 입사자 장비 신청 절차 알려줘", "IT장비신청절차.txt"),
        ("IT", "보안 교육 언제까지 이수해야 해?", "보안교육이수규정.txt"),
        ("검증용 (제한 답변)", "회사 주식 상장 일정이 언제야?", "→ \"확인할 수 없습니다\" 출력"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  9장: 성능 테스트
# ══════════════════════════════════════════════════════════════

heading1("성능 테스트 결과", "9️⃣")

heading2("인덱싱 성능", "📊")
make_table(
    ["측정 항목", "결과"],
    [
        ("처리한 총 파일 수", "30개 (PDF 20개 + TXT 10개)"),
        ("파싱 성공률", "100%  (30/30)"),
        ("생성된 chunk 수", "67개"),
        ("임베딩 + DB 저장 소요 시간", "23.2초"),
        ("저장된 Vector DB 크기", "2.5 MB"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("RAG 질의응답 품질 테스트", "📊")
make_table(
    ["질문", "근거 있음", "유사도 점수", "응답 시간"],
    [
        ("연차 신청 기한",           "✅ 있음", "0.727", "31초 (첫 질문, 모델 로딩 포함)"),
        ("출장비 영수증 제출 방법",  "✅ 있음", "0.705", "6초"),
        ("복지포인트 잔액 확인",     "✅ 있음", "0.788", "7초"),
        ("신규 입사자 장비 신청",    "✅ 있음", "0.762", "9초"),
        ("경조사 지원금 신청",       "✅ 있음", "0.796", "5초"),
        ("회사 주식 상장 일정 (검증)", "✅ 제한 답변", "0.510", "6초"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("유사도 점수 기준표", "📏")
make_table(
    ["점수 범위", "의미", "시스템 동작"],
    [
        ("0.70 이상", "매우 관련성 높음", "✅ 문서 근거 답변 생성"),
        ("0.50 ~ 0.69", "관련성 있음", "✅ 문서 근거 답변 생성"),
        ("0.30 ~ 0.49", "낮은 관련성", "⚠️ 문서 근거 답변 (주의 필요)"),
        ("0.30 미만", "관련 문서 없음", "🚫 \"확인할 수 없습니다\" 출력"),
    ],
    header_bg=BLUE_DARK,
    alt_bg=BLUE_LIGHT,
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  10장: 확장 로드맵 + 용어 정리
# ══════════════════════════════════════════════════════════════

heading1("다음 확장 방향 및 용어 정리", "🔟")

heading2("MVP → 운영 시스템 로드맵", "🗺️")
make_table(
    ["단계", "현재 (POC)", "다음 (MVP)", "이후 (운영)"],
    [
        ("UI", "Streamlit 단일 앱", "React/Next.js", "모바일 앱 포함"),
        ("백엔드", "Python 단일 앱", "FastAPI REST API", "마이크로서비스"),
        ("Vector DB", "Chroma 로컬", "Qdrant 고성능 DB", "분산 클러스터"),
        ("LLM", "Ollama 로컬", "vLLM 고성능 서버", "GPU 서버 운영"),
        ("문서 수", "30개 샘플", "100개+ 실문서", "전사 문서 연동"),
        ("인증", "없음", "기초 사용자 권한", "LDAP/SSO 연동"),
        ("문서 권한", "없음", "부서별 기초 필터", "세분화 권한 제어"),
        ("배포", "로컬 PC", "사내 서버", "Docker·K8s 운영"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

heading2("추가 가능한 기능들", "➕")
make_table(
    ["기능", "설명", "난이도"],
    [
        ("문서 권한 필터", "부서별로 볼 수 있는 문서 구분", "중"),
        ("OCR 지원", "스캔된 이미지 PDF도 텍스트 추출", "중"),
        ("파일 서버 연동", "SharePoint·Google Drive 자동 동기화", "중"),
        ("관리자 대시보드", "인덱싱 현황·자주 묻는 질문 통계", "중"),
        ("다국어 지원", "영어 문서도 한국어로 답변", "중"),
        ("슬랙/Teams 연동", "메신저에서 바로 질문", "중"),
        ("모바일 앱", "스마트폰에서 질문", "고"),
        ("감사 로그", "누가 언제 무엇을 질문했는지 기록", "저"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

divider()
heading2("용어 정리 (Glossary)", "📖")
make_table(
    ["용어", "설명"],
    [
        ("RAG", "Retrieval-Augmented Generation의 약자.\n\"검색-보강-생성\" 3단계로 동작하는 AI 기술"),
        ("LLM", "Large Language Model(대형 언어 모델).\nChatGPT 같은 대화형 AI의 기반 기술"),
        ("임베딩 (Embedding)", "텍스트를 숫자 벡터로 변환하는 기술.\n의미가 비슷한 문장은 비슷한 숫자를 가짐"),
        ("벡터 DB (Vector DB)", "숫자(벡터)를 저장하고 유사도로 검색하는 DB.\n이 프로젝트에서는 Chroma 사용"),
        ("Ollama", "오픈소스 로컬 AI 실행 플랫폼.\n인터넷 없이 PC에서 LLM을 실행할 수 있게 해줌"),
        ("Streamlit", "Python으로 웹 앱을 빠르게 만드는 프레임워크"),
        ("청크 (Chunk)", "긴 문서를 작게 조각낸 단위. 이 프로젝트는 1,600자 기준"),
        ("바이브코딩 (Vibe Coding)", "자연어로 AI에게 지시해 코드를 생성하는 개발 방식"),
        ("On-Premise (온프레미스)", "인터넷 없이 자체 서버·PC에서만 운영하는 방식"),
        ("POC", "Proof of Concept. 개념 검증용 시범 구현"),
        ("코사인 유사도", "두 벡터 사이의 각도로 유사성을 측정하는 방법.\n값이 1에 가까울수록 유사"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

divider()
heading2("참고 자료", "🔗")
make_table(
    ["항목", "주소"],
    [
        ("GitHub 소스코드", "https://github.com/CharlyChoiStory/local-enterprise-rag-chatbot"),
        ("Ollama 공식 사이트", "https://ollama.com"),
        ("Streamlit 공식 문서", "https://docs.streamlit.io"),
        ("Chroma DB", "https://www.trychroma.com"),
        ("llama3.1 모델", "https://ollama.com/library/llama3.1"),
        ("bge-m3 임베딩 모델", "https://ollama.com/library/bge-m3"),
    ],
    header_bg=BLUE_DARK, alt_bg=BLUE_LIGHT,
)

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(
    "이 매뉴얼은 AI 도구(텔레그램 → OpenClaw 에이전트 + Claude Code)를 활용한 바이브코딩으로 제작되었습니다.\n"
    "개발 환경: VSCode + Claude Code  |  원격 지시: 텔레그램 → OpenClaw  |  코드 공유: GitHub"
)
r_f.font.size  = Pt(9)
r_f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
r_f.font.italic = True

# ── 저장 ─────────────────────────────────────────────────────
output = "MANUAL.docx"
doc.save(output)
print(f"✅  {output} 생성 완료")
